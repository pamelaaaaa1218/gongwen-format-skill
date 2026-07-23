#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
标准公文格式套用脚本
将 Word 文档（.docx / .doc）排版为国企标准化公文格式。

用法:
    python3 apply_gongwen_format.py 输入文件 [-o 输出文件] [--title-lines N] [--report-only]
                                    [--keep-numbering] [--keep-quotes] [--keep-spaces]

参数:
    输入文件          待排版的 .docx 或 .doc 文件
    -o 输出文件       输出路径（默认在原文件旁生成「原名（公文格式）.docx」）
    --title-lines N   强制把开头 N 个非空段落识别为「标题区」（自动识别不准时用）
    --report-only     只打印段落分类，不生成文件（用于先预览识别是否正确）
    --keep-numbering  保留原编号（不做 1.1→（一）、（1）→1. 的规范化）
    --keep-quotes     保留英文双引号，不转中文 “ ”
    --keep-spaces     保留数字/英文两侧的空格，不做清理
"""
import sys
import os
import re
import argparse
import shutil
import subprocess
import tempfile

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("缺少 python-docx，请先运行:  pip3 install python-docx")


# ============================================================
#  标准公文格式规格 —— 如需调整格式，改这里的值即可
#  （遵循《党政机关公文格式 GB/T 9704-2012》，标题采用集团惯用的小二号）
# ============================================================
FONT_TITLE = "方正小标宋简体"   # 标题 / 单位抬头
FONT_BODY  = "仿宋_GB2312"      # 正文、三级「1.」/四级「（1）」序号、附件、落款
FONT_H1    = "黑体"             # 一级标题「一、」
FONT_H2    = "楷体_GB2312"      # 二级标题「（一）」
FONT_WESTERN = "Times New Roman"  # 数字与英文（西文）字体；中文仍用上面的东亚字体

SIZE_TITLE = 18    # 标题字号：小二号 = 18pt
SIZE_BODY  = 16    # 正文字号：三号  = 16pt

LINE_SPACING_PT  = 28   # 行距：固定值 28 磅
FIRST_LINE_CHARS = 2    # 正文/标题序号首行缩进：2 字符
H1_BEFORE_LINES  = 1.0  # 一级标题段前空 1 行
H2_BEFORE_LINES  = 0.5  # 二级标题段前空 0.5 行
TITLE_AFTER_LINES = 1.5  # 大标题段后空 1.5 行
TABLE_BEFORE_LINES = 0.5  # 表格整体段前空 0.5 行（落在表格上方相邻段落的段后）

# 页面设置（A4，单位：厘米）
PAGE_WIDTH_CM, PAGE_HEIGHT_CM = 21.0, 29.7
MARGIN_TOP_CM, MARGIN_BOTTOM_CM = 3.7, 3.5
MARGIN_LEFT_CM, MARGIN_RIGHT_CM = 2.8, 2.6
# ============================================================


# ---- 段落类型识别规则 ----
RE_H1     = re.compile(r'^\s*[一二三四五六七八九十]+[、.．]')
# 二级标题：标准公文的「（一）」，或 AI 常写的数字分级「1.1」「2.3」
# 数字式要求编号后紧跟分隔符（空格/顿号/右括号/行尾），以免误伤「1.1万元」这类正文
RE_H2CN   = re.compile(r'^\s*[（(]\s*[一二三四五六七八九十]+\s*[）)]')
RE_H2NUM  = re.compile(r'^\s*\d+\.\d+(?=[ \t　、．.）)]|$)')
RE_H2CAP  = re.compile(r'^\s*(\d+)\.(\d+)[ \t　、．.]*\s*(.*)$', re.S)
RE_H2     = re.compile(r'^\s*(?:[（(]\s*[一二三四五六七八九十]+\s*[）)]'
                       r'|\d+\.\d+(?=[ \t　、．.）)]|$))')
# 三级序号「1.」（排除「1.1」数字分级——那是二级）；四级序号「（1）」
RE_L3     = re.compile(r'^\s*\d+\s*[、.．](?!\d)')
RE_L4     = re.compile(r'^\s*[（(]\s*(\d+)\s*[）)]\s*(.*)$', re.S)
RE_TITLE  = re.compile(r'关于.{0,80}的.{0,40}'
                       r'(报告|汇报|通知|请示|意见|方案|总结|计划|纪要|函|批复|决定|通报|公告|说明|情况)')
RE_DATE   = re.compile(r'^[\s\d〇○零一二三四五六七八九十廿卅×*]{2,14}年'
                       r'[\s\d〇○零一二三四五六七八九十廿卅×*]{1,5}月'
                       r'[\s\d〇○零一二三四五六七八九十廿卅×*]{1,5}日\s*$')
RE_ATTACH = re.compile(r'^\s*附\s*件\s*\d*\s*[:：]?')

ROLE_CN = {
    'title': '标题', 'salutation': '称谓', 'h1': '一级标题', 'h2': '二级标题',
    'attachment': '附件', 'sign_unit': '落款单位', 'sign_date': '落款日期',
    'body': '正文', 'empty': '空行',
}
ROLE_FONT = {
    'title': (FONT_TITLE, SIZE_TITLE),
    'salutation': (FONT_BODY, SIZE_BODY),
    'h1': (FONT_H1, SIZE_BODY),
    'h2': (FONT_H2, SIZE_BODY),
    'attachment': (FONT_BODY, SIZE_BODY),
    'sign_unit': (FONT_BODY, SIZE_BODY),
    'sign_date': (FONT_BODY, SIZE_BODY),
    'body': (FONT_BODY, SIZE_BODY),
}


def is_salutation(text):
    """称谓行：短，且以冒号结尾，如「各位领导：」"""
    return 0 < len(text) <= 14 and text.endswith(('：', ':'))


def classify(texts, title_lines=None):
    """根据段落文字判断每段的角色，返回与 texts 等长的角色列表。"""
    n = len(texts)
    roles = [None] * n
    nonempty = [i for i in range(n) if texts[i]]
    if not nonempty:
        return ['empty'] * n

    # ---- 标题区 ----
    title_idxs = []
    if title_lines and title_lines > 0:
        title_idxs = nonempty[:title_lines]
    else:
        for i in nonempty:
            t = texts[i]
            if RE_TITLE.search(t):
                title_idxs.append(i)
                break
            if RE_H1.match(t) or RE_H2.match(t) or RE_ATTACH.match(t) or is_salutation(t):
                break
            if len(t) > 45:                       # 过长，更像正文
                if not title_idxs:
                    title_idxs.append(i)          # 但若还没有标题，则视为长标题
                break
            title_idxs.append(i)
            if len(title_idxs) >= 3:
                break
        if not title_idxs:
            title_idxs = [nonempty[0]]
    for i in title_idxs:
        roles[i] = 'title'
    last_title = max(title_idxs)

    # ---- 落款（单位 + 日期）----
    sign_date_idx = None
    for i in reversed(nonempty):
        if i <= last_title:
            break
        if RE_DATE.match(texts[i]):
            sign_date_idx = i
            break
    if sign_date_idx is not None:
        roles[sign_date_idx] = 'sign_date'
        for i in reversed([j for j in nonempty if last_title < j < sign_date_idx]):
            if len(texts[i]) <= 25 and not RE_H1.match(texts[i]) and not RE_H2.match(texts[i]):
                roles[i] = 'sign_unit'
            break

    # ---- 称谓（标题之后第一段）----
    after = [i for i in nonempty if i > last_title]
    if after and roles[after[0]] is None and is_salutation(texts[after[0]]):
        roles[after[0]] = 'salutation'

    # ---- 其余段落 ----
    for i in range(n):
        if roles[i] is not None:
            continue
        t = texts[i]
        if not t:
            roles[i] = 'empty'
        elif RE_H1.match(t):
            roles[i] = 'h1'
        elif RE_H2.match(t):
            roles[i] = 'h2'
        elif RE_ATTACH.match(t):
            roles[i] = 'attachment'
        else:
            roles[i] = 'body'
    return roles


def set_run(run, font_name, size_pt):
    """统一一个文字 run 的字体、字号，并清除多余的颜色/加粗。"""
    run.font.size = Pt(size_pt)
    run.font.bold = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.insert(0, rfonts)
    # 中文（东亚）用 font_name，数字/英文（西文）用 Times New Roman
    rfonts.set(qn('w:eastAsia'), font_name)
    rfonts.set(qn('w:ascii'), FONT_WESTERN)
    rfonts.set(qn('w:hAnsi'), FONT_WESTERN)
    rfonts.set(qn('w:cs'), FONT_WESTERN)
    if rfonts.get(qn('w:hint')) is not None:
        del rfonts.attrib[qn('w:hint')]
    for attr in ('w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme', 'w:cstheme'):
        if rfonts.get(qn(attr)) is not None:
            del rfonts.attrib[qn(attr)]


def set_run_font_only(run, font_name):
    """仅统一字体（用于表格单元格，保留原字号/加粗）。中文用 font_name，数字/英文用西文字体。"""
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.insert(0, rfonts)
    rfonts.set(qn('w:eastAsia'), font_name)
    rfonts.set(qn('w:ascii'), FONT_WESTERN)
    rfonts.set(qn('w:hAnsi'), FONT_WESTERN)
    rfonts.set(qn('w:cs'), FONT_WESTERN)
    if rfonts.get(qn('w:hint')) is not None:
        del rfonts.attrib[qn('w:hint')]


def _set_before(spacing, pt, before_lines):
    """在 w:spacing 上写段前：beforeLines(1/100 行) + before(twip 兜底)。"""
    if before_lines:
        spacing.set(qn('w:beforeLines'), str(int(before_lines * 100)))
        spacing.set(qn('w:before'), str(int(pt * 20 * before_lines)))
    else:
        spacing.set(qn('w:before'), '0')
        if spacing.get(qn('w:beforeLines')) is not None:
            del spacing.attrib[qn('w:beforeLines')]


def _set_after(spacing, pt, after_lines):
    """在 w:spacing 上写段后：afterLines(1/100 行) + after(twip 兜底)。"""
    if after_lines:
        spacing.set(qn('w:afterLines'), str(int(after_lines * 100)))
        spacing.set(qn('w:after'), str(int(pt * 20 * after_lines)))
    else:
        spacing.set(qn('w:after'), '0')
        if spacing.get(qn('w:afterLines')) is not None:
            del spacing.attrib[qn('w:afterLines')]


def set_line_spacing(p, pt, before_lines=0, after_lines=0):
    """固定值行距；段前 before_lines 行、段后 after_lines 行（0 = 无）。
    「N 行」用 beforeLines/afterLines（Word/WPS 里显示为「N 行」），再按行距折算 twip 兜底。"""
    pPr = p._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), str(int(pt * 20)))
    spacing.set(qn('w:lineRule'), 'exact')
    for a in ('w:beforeAutospacing', 'w:afterAutospacing'):
        if spacing.get(qn(a)) is not None:
            del spacing.attrib[qn(a)]
    _set_before(spacing, pt, before_lines)
    _set_after(spacing, pt, after_lines)


def set_para_after(p_element, pt, after_lines):
    """只设置某段落 element 的段后（不动其行距/段前），用于表格上方段落。"""
    pPr = p_element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_element.insert(0, pPr)
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    if spacing.get(qn('w:afterAutospacing')) is not None:
        del spacing.attrib[qn('w:afterAutospacing')]
    _set_after(spacing, pt, after_lines)


def apply_table_spacing(doc, before_lines=TABLE_BEFORE_LINES):
    """表格整体段前 N 行：落在表格上方相邻段落的段后。
    表格若在文首或紧接另一表格（上方无段落）则跳过。天然幂等——一旦上方存在段落即只改其段后。"""
    body = doc.element.body
    n = 0
    for tbl in body.findall(qn('w:tbl')):
        prev = tbl.getprevious()
        if prev is not None and prev.tag == qn('w:p'):
            set_para_after(prev, LINE_SPACING_PT, before_lines)
            n += 1
    return n


def set_first_line_indent(p, chars):
    """首行缩进 chars 字符（chars=0 表示顶格）。"""
    pPr = p._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    for a in ('w:firstLine', 'w:firstLineChars', 'w:hanging', 'w:hangingChars',
              'w:left', 'w:leftChars', 'w:start', 'w:startChars'):
        if ind.get(qn(a)) is not None:
            del ind.attrib[qn(a)]
    if chars:
        ind.set(qn('w:firstLineChars'), str(int(chars * 100)))
        ind.set(qn('w:firstLine'), str(int(chars * SIZE_BODY * 20)))


def cn_num(n):
    """阿拉伯数字转公文中文序号：1→一、10→十、11→十一、21→二十一。"""
    digits = '〇一二三四五六七八九'
    if n <= 0:
        return str(n)
    if n < 10:
        return digits[n]
    if n == 10:
        return '十'
    if n < 20:
        return '十' + digits[n % 10]
    if n < 100:
        t, o = divmod(n, 10)
        return digits[t] + '十' + (digits[o] if o else '')
    return str(n)


def set_para_text(p, text):
    """把整段文字重写为 text，保留首个 run 的格式（用于改写标题段的编号）。"""
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(text)


def plan_numbering(roles, texts):
    """规划二级标题编号规范化：数字分级「1.1」→ 公文括号序号「（一）」。
    只处理被识别为二级标题的段落；已是「（一）」式的跳过；序号取小数点后的数字，
    因此每个一级标题下会自动从「（一）」重新计数。返回 {段索引: 新文字}。"""
    plan = {}
    for i, t in enumerate(texts):
        if roles[i] != 'h2' or RE_H2CN.match(t):
            continue
        m = RE_H2CAP.match(t)
        if m:
            plan[i] = '（%s）%s' % (cn_num(int(m.group(2))), m.group(3).strip())
    return plan


def apply_numbering(doc, plan):
    for i, new_text in plan.items():
        set_para_text(doc.paragraphs[i], new_text)


def plan_l3_promotion(roles, texts):
    """三级优先：公文层级为 一、→（一）→ 1. →（1），不允许跳过三级直接用四级。
    若某节（上一个标题之后）尚未出现过三级序号「1.」，却直接出现「（1）（2）」，
    则把该节里这些段落提升为三级「1. 2.」（保留原数字）。
    一旦该节先出现了真正的三级序号，其后的「（N）」视为正常四级，保持不变。"""
    plan = {}
    l3_seen = False
    for i, t in enumerate(texts):
        if roles[i] in ('title', 'h1', 'h2'):
            l3_seen = False          # 新的一节，重新计
            continue
        if roles[i] != 'body' or not t:
            continue
        if RE_L3.match(t):
            l3_seen = True           # 本节已有三级序号
            continue
        m = RE_L4.match(t)
        if m and not l3_seen:
            plan[i] = '%s.%s' % (m.group(1), m.group(2).strip())
    return plan


def _convert_quotes_in_para(p):
    """把一个段落里的英文直双引号 " 按段内出现次序交替转成中文「“」「”」。
    逐 run 处理、不破坏 run 结构；每段独立配对（引号极少跨段）。"""
    opening = True
    for run in p.runs:
        if '"' not in run.text:
            continue
        out = []
        for ch in run.text:
            if ch == '"':
                out.append('“' if opening else '”')
                opening = not opening
            else:
                out.append(ch)
        run.text = ''.join(out)


def apply_quotes(doc):
    """全文（含表格单元格）英文双引号 → 中文双引号。"""
    for p in doc.paragraphs:
        _convert_quotes_in_para(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _convert_quotes_in_para(p)


def count_quotes(doc):
    n = 0
    for p in doc.paragraphs:
        n += p.text.count('"')
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    n += p.text.count('"')
    return n


# ---- 数字/英文两侧多余空格清理 ----
# 中文与全角字符（含全角标点、中文引号、破折号、间隔号等）
RE_CJK_CH = re.compile(r'[\u2e80-\u303f\u3400-\u4dbf\u4e00-\u9fff\uf900-\ufaff'
                       r'\ufe30-\ufe4f\uff00-\uffef\u2014\u2018\u2019\u201c\u201d\u2026\u00b7]')
RE_GAP    = re.compile(r'(\S)[ \t\u00a0\u2000-\u200a\u202f\u3000]+(\S)')
_SPACE_CHARS = (' \t\u00a0\u3000\u2000\u2001\u2002\u2003\u2004\u2005'
                '\u2006\u2007\u2008\u2009\u200a\u202f')


def _gap_decide(m):
    """决定一处空格去留：
    任一侧是中文/全角 → 删（「共 23 条」「用 Python 写」）；
    两侧都是字母/数字   → 留（「Claude Code」「GB/T 9704」）；
    数字挨着符号       → 删（「2514 / 点赞」→「2514/点赞」）；
    其余（纯英文语境的标点）→ 留（「e.g. test」）。"""
    L, R = m.group(1), m.group(2)
    if RE_CJK_CH.match(L) or RE_CJK_CH.match(R):
        return L + R
    if L.isalnum() and R.isalnum():
        return m.group(0)
    if L.isdigit() or R.isdigit():
        return L + R
    return m.group(0)


def clean_spaces_text(s):
    """对一段纯文本做空格清理（循环到不再变化，处理连续边界）。"""
    prev = None
    while prev != s:
        prev = s
        s = RE_GAP.sub(_gap_decide, s)
    return s


def _clean_spaces_runs(runs):
    """逐 run 清理空格，不合并 run（保住图片/域代码/换行等特殊元素）。
    用相邻 run 的边界字符做上下文，跨 run 的空格也能正确判断。"""
    texts = [r.text for r in runs]
    removed = 0
    for i, t in enumerate(texts):
        if not t:
            continue
        prev_ch = next_ch = ''
        for j in range(i - 1, -1, -1):
            st = texts[j].rstrip(_SPACE_CHARS)
            if st:
                prev_ch = st[-1]
                break
        for j in range(i + 1, len(texts)):
            st = texts[j].lstrip(_SPACE_CHARS)
            if st:
                next_ch = st[0]
                break
        padded = prev_ch + t + next_ch
        cleaned = clean_spaces_text(padded)
        if prev_ch:
            cleaned = cleaned[1:]
        if next_ch:
            cleaned = cleaned[:-1]
        if cleaned != t:
            removed += len(t) - len(cleaned)
            texts[i] = cleaned
    for r, t in zip(runs, texts):
        if r.text != t:
            r.text = t
    return removed


def apply_spaces(doc):
    """全文（含表格单元格）清理数字/英文两侧多余空格，返回删除的空格数。"""
    removed = 0
    for p in doc.paragraphs:
        removed += _clean_spaces_runs(p.runs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    removed += _clean_spaces_runs(p.runs)
    return removed


def count_spaces(doc):
    """统计将被清理的空格数（预览用）。"""
    n = 0
    for p in doc.paragraphs:
        n += len(p.text) - len(clean_spaces_text(p.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    n += len(p.text) - len(clean_spaces_text(p.text))
    return n


def format_document(doc, roles, texts):
    """按角色把格式套到每个段落上。"""
    counts = {}
    for i, p in enumerate(doc.paragraphs):
        role = roles[i]
        counts[role] = counts.get(role, 0) + 1
        before = H1_BEFORE_LINES if role == 'h1' else (
                 H2_BEFORE_LINES if role == 'h2' else 0)
        after = TITLE_AFTER_LINES if role == 'title' else 0
        set_line_spacing(p, LINE_SPACING_PT, before_lines=before, after_lines=after)
        if role == 'empty':
            continue
        font_name, size = ROLE_FONT[role]
        for run in p.runs:
            set_run(run, font_name, size)
        if role == 'title':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_first_line_indent(p, 0)
        elif role in ('sign_unit', 'sign_date'):
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_first_line_indent(p, 0)
        elif role in ('salutation', 'attachment'):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_first_line_indent(p, 0)
        else:  # body / h1 / h2
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_first_line_indent(p, FIRST_LINE_CHARS)

    # 页面设置
    for section in doc.sections:
        section.page_width = Cm(PAGE_WIDTH_CM)
        section.page_height = Cm(PAGE_HEIGHT_CM)
        section.top_margin = Cm(MARGIN_TOP_CM)
        section.bottom_margin = Cm(MARGIN_BOTTOM_CM)
        section.left_margin = Cm(MARGIN_LEFT_CM)
        section.right_margin = Cm(MARGIN_RIGHT_CM)

    # 表格：仅统一字体，不改字号/结构
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font_only(run, FONT_BODY)
    return counts


def print_report(texts, roles, n_tables, num_plan=None, l3_plan=None,
                 quote_count=0, space_count=0):
    print('=' * 60)
    print('  标准公文格式 —— 段落分类结果')
    print('=' * 60)
    print(f"{'序号':<5}{'类型':<10}内容预览")
    print('-' * 60)
    for i, role in enumerate(roles):
        if role == 'empty':
            continue
        preview = texts[i][:30].replace('\n', ' ')
        print(f"{i:<6}{ROLE_CN.get(role, role):<11}{preview}")
    print('-' * 60)
    summary = '  '.join(f"{ROLE_CN[r]}{c}"
                         for r, c in sorted(
                             {r: roles.count(r) for r in set(roles) if r != 'empty'}.items()))
    print('统计：' + summary + (f"  表格{n_tables}个" if n_tables else ''))
    if num_plan:
        print('-' * 60)
        print('二级编号规范化（数字分级 → 公文括号序号）：')
        for i in sorted(num_plan):
            print(f"  {texts[i][:22]:<24} →  {num_plan[i][:22]}")
    if l3_plan:
        print('-' * 60)
        print('三级优先（该节无「1.」三级序号，「（1）」式提升为三级）：')
        for i in sorted(l3_plan):
            print(f"  {texts[i][:22]:<24} →  {l3_plan[i][:22]}")
    if quote_count:
        print('-' * 60)
        print(f'标点规范化：{quote_count} 处英文双引号 \" 将转为中文 “ ”')
    if space_count:
        print('-' * 60)
        print(f'空格规范化：将删除数字/英文两侧多余空格 {space_count} 处')
    print('-' * 60)
    print('若识别有误：标题行数不对用 --title-lines N 重跑；个别段落错位可手动微调。')
    print('规范化默认开启，保留原样用 --keep-numbering / --keep-quotes / --keep-spaces。')
    print('=' * 60)


def main():
    ap = argparse.ArgumentParser(description='标准公文格式套用工具')
    ap.add_argument('input', help='输入文件 .docx 或 .doc')
    ap.add_argument('-o', '--output', help='输出文件路径')
    ap.add_argument('--title-lines', type=int, help='强制把开头 N 个非空段落作为标题区')
    ap.add_argument('--report-only', action='store_true', help='只打印分类，不生成文件')
    ap.add_argument('--keep-numbering', action='store_true',
                    help='保留原编号，不把数字分级「1.1」转成公文序号「（一）」')
    ap.add_argument('--keep-quotes', action='store_true',
                    help='保留英文引号，不转中文引号')
    ap.add_argument('--keep-spaces', action='store_true',
                    help='保留数字/英文两侧的空格，不做清理')
    args = ap.parse_args()

    src = os.path.abspath(os.path.expanduser(args.input))
    if not os.path.isfile(src):
        sys.exit(f"找不到文件：{src}")

    # .doc 先转换为 .docx
    tmpdir = None
    work_path = src
    if src.lower().endswith('.doc') and not src.lower().endswith('.docx'):
        if not shutil.which('textutil'):
            sys.exit("检测到 .doc 文件，但本机没有 textutil（仅 macOS 自带）。\n"
                     "请先在 Word/WPS 里把文件「另存为」.docx 后再运行。")
        tmpdir = tempfile.mkdtemp()
        work_path = os.path.join(tmpdir, 'converted.docx')
        try:
            subprocess.run(['textutil', '-convert', 'docx', src, '-output', work_path],
                           check=True, capture_output=True)
        except subprocess.CalledProcessError as e:
            sys.exit(f".doc 转换失败：{e.stderr.decode('utf-8', 'ignore')}")
        print(f"已将 .doc 转换为 .docx 进行处理：{os.path.basename(src)}")

    try:
        doc = Document(work_path)
    except Exception as e:
        sys.exit(f"无法打开文档（可能不是有效的 Word 文件）：{e}")

    texts = [p.text.strip() for p in doc.paragraphs]
    if not any(texts):
        sys.exit("文档中没有可识别的文字段落。")

    roles = classify(texts, title_lines=args.title_lines)
    num_plan = {} if args.keep_numbering else plan_numbering(roles, texts)
    l3_plan = {} if args.keep_numbering else plan_l3_promotion(roles, texts)
    quote_count = 0 if args.keep_quotes else count_quotes(doc)
    space_count = 0 if args.keep_spaces else count_spaces(doc)
    print_report(texts, roles, len(doc.tables), num_plan, l3_plan,
                 quote_count, space_count)

    if args.report_only:
        if tmpdir:
            shutil.rmtree(tmpdir, ignore_errors=True)
        return

    if num_plan:
        apply_numbering(doc, num_plan)
    if l3_plan:
        apply_numbering(doc, l3_plan)
    if not args.keep_quotes:
        apply_quotes(doc)
    removed_spaces = 0 if args.keep_spaces else apply_spaces(doc)
    counts = format_document(doc, roles, texts)
    table_gaps = apply_table_spacing(doc)  # 须在 format_document 之后，避免段后被清零覆盖

    if args.output:
        out = os.path.abspath(os.path.expanduser(args.output))
    else:
        stem = os.path.splitext(os.path.basename(src))[0]
        out = os.path.join(os.path.dirname(src), f"{stem}（公文格式）.docx")

    doc.save(out)
    if tmpdir:
        shutil.rmtree(tmpdir, ignore_errors=True)

    print()
    print(f"✅ 已生成：{out}")
    print("   字体/字号、行距(固定28磅)、首行缩进、页边距(A4·上3.7/下3.5/左2.8/右2.6cm)均已套用。")
    print("   数字与英文已统一为 Times New Roman，中文保持仿宋/黑体/楷体。")
    print("   间距：大标题段后1.5行、一级标题段前1行、二级标题段前0.5行。")
    if table_gaps:
        print(f"   表格：{table_gaps} 处表格已设段前0.5行（落在表格上方段落的段后）。")
    if num_plan:
        print(f"   二级编号：{len(num_plan)} 处数字分级「1.1」已转为公文序号「（一）」。")
    if l3_plan:
        print(f"   三级优先：{len(l3_plan)} 处「（1）」式序号已提升为三级「1.」式。")
    if quote_count and not args.keep_quotes:
        print(f"   标点：{quote_count} 处英文双引号已转为中文「“ ”」。")
    if removed_spaces:
        print(f"   空格：已删除数字/英文两侧多余空格 {removed_spaces} 处。")
    print("   注：文字颜色已统一为黑色、清除了多余加粗；表格仅统一字体未改字号。")


if __name__ == '__main__':
    main()
